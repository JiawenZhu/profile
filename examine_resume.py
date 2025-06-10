#!/usr/bin/env python3

from docx import Document
import sys

def examine_resume(filename):
    """Examine the resume content to understand what needs to be changed"""
    
    try:
        doc = Document(filename)
        
        print(f"Examining resume: {filename}")
        print("=" * 50)
        
        all_text = []
        
        # Extract text from main document
        print("\n📄 MAIN DOCUMENT CONTENT:")
        print("-" * 30)
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                text = paragraph.text.strip()
                all_text.append(text)
                print(f"{i+1:2d}: {text}")
        
        # Extract text from tables
        print("\n📊 TABLE CONTENT:")
        print("-" * 30)
        table_num = 0
        for table in doc.tables:
            table_num += 1
            print(f"\nTable {table_num}:")
            for row_num, row in enumerate(table.rows):
                for col_num, cell in enumerate(row.cells):
                    if cell.text.strip():
                        text = cell.text.strip()
                        all_text.append(text)
                        print(f"  Row {row_num+1}, Col {col_num+1}: {text}")
        
        # Extract text from headers and footers
        print("\n📋 HEADERS & FOOTERS:")
        print("-" * 30)
        for section_num, section in enumerate(doc.sections):
            print(f"\nSection {section_num + 1}:")
            
            # Header
            header = section.header
            if header:
                for i, paragraph in enumerate(header.paragraphs):
                    if paragraph.text.strip():
                        text = paragraph.text.strip()
                        all_text.append(text)
                        print(f"  Header {i+1}: {text}")
            
            # Footer
            footer = section.footer
            if footer:
                for i, paragraph in enumerate(footer.paragraphs):
                    if paragraph.text.strip():
                        text = paragraph.text.strip()
                        all_text.append(text)
                        print(f"  Footer {i+1}: {text}")
        
        # Search for specific patterns
        print("\n🔍 PATTERN ANALYSIS:")
        print("-" * 30)
        
        full_text = " ".join(all_text).lower()
        
        # Check for names
        if "evan" in full_text:
            print("✅ Found 'Evan' in document")
        else:
            print("❌ 'Evan' not found in document")
        
        if "jiawen" in full_text:
            print("✅ Found 'Jiawen' in document")
        else:
            print("❌ 'Jiawen' not found in document")
        
        # Check for URLs and project references
        url_patterns = [
            "github.com",
            "churchinurbana",
            "healthwell",
            "liftoff",
            "vercel",
            "http",
            "https",
            ".com",
            ".org"
        ]
        
        found_urls = []
        for pattern in url_patterns:
            if pattern in full_text:
                found_urls.append(pattern)
        
        if found_urls:
            print(f"✅ Found URL patterns: {', '.join(found_urls)}")
        else:
            print("❌ No URL patterns found")
        
        # Look for project names
        project_patterns = [
            "github profile",
            "church",
            "quiz",
            "healthwell",
            "liftoff"
        ]
        
        found_projects = []
        for pattern in project_patterns:
            if pattern in full_text:
                found_projects.append(pattern)
        
        if found_projects:
            print(f"✅ Found project patterns: {', '.join(found_projects)}")
        else:
            print("❌ No project patterns found")
        
        # Show lines containing key terms
        print("\n🎯 LINES WITH KEY TERMS:")
        print("-" * 30)
        
        key_terms = ["evan", "github", "church", "http", "project"]
        for term in key_terms:
            matching_lines = [line for line in all_text if term.lower() in line.lower()]
            if matching_lines:
                print(f"\n'{term}' found in:")
                for line in matching_lines[:3]:  # Show first 3 matches
                    print(f"  → {line}")
                if len(matching_lines) > 3:
                    print(f"  ... and {len(matching_lines) - 3} more")
        
        return True
        
    except Exception as e:
        print(f"Error examining resume: {str(e)}")
        return False

def main():
    resume_file = "Jiawen-Resume-2025.docx"
    examine_resume(resume_file)

if __name__ == "__main__":
    main() 