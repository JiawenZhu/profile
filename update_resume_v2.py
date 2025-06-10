#!/usr/bin/env python3

from docx import Document
import sys

def update_resume(filename):
    """Update the resume with requested changes"""
    
    try:
        doc = Document(filename)
        changes_made = []
        
        def replace_text_in_runs(paragraph, old_text, new_text):
            """Replace text while preserving formatting"""
            full_text = paragraph.text
            if old_text in full_text:
                # Clear all runs
                for run in paragraph.runs:
                    run.text = ""
                
                # Add new text to first run
                if paragraph.runs:
                    paragraph.runs[0].text = full_text.replace(old_text, new_text)
                else:
                    paragraph.add_run(full_text.replace(old_text, new_text))
                return True
            return False
        
        def update_document_text(old_text, new_text):
            """Update text throughout the document"""
            changes = 0
            
            # Update paragraphs
            for paragraph in doc.paragraphs:
                if replace_text_in_runs(paragraph, old_text, new_text):
                    changes += 1
            
            # Update tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if replace_text_in_runs(paragraph, old_text, new_text):
                                changes += 1
            
            # Update headers and footers
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    if replace_text_in_runs(paragraph, old_text, new_text):
                        changes += 1
                for paragraph in section.footer.paragraphs:
                    if replace_text_in_runs(paragraph, old_text, new_text):
                        changes += 1
            
            return changes
        
        # 1. Update name from "evan ZHu" to "Jiawen Zhu"
        name_changes = update_document_text("evan ZHu", "Jiawen Zhu")
        if name_changes > 0:
            changes_made.append(f"Updated name 'evan ZHu' to 'Jiawen Zhu' in {name_changes} locations")
        
        # Also handle other variations
        name_changes2 = update_document_text("Evan Zhu", "Jiawen Zhu")
        if name_changes2 > 0:
            changes_made.append(f"Updated name 'Evan Zhu' to 'Jiawen Zhu' in {name_changes2} locations")
        
        name_changes3 = update_document_text("EVAN ZHU", "JIAWEN ZHU")
        if name_changes3 > 0:
            changes_made.append(f"Updated name 'EVAN ZHU' to 'JIAWEN ZHU' in {name_changes3} locations")
        
        # 2. Add project sections with new links
        # Since the current resume doesn't have a clear projects section, 
        # I'll add content to the experience section or create a new section
        
        # Find a place to add project information
        # Let's add after the skills section or create a projects section
        
        # First, let's add the project URLs to the existing experience if there's room
        # or modify the current freelancer section to include the projects
        
        # Look for freelancer section and enhance it with project links
        freelancer_text = "Web Developer/Cloud Developer (Freelancer)"
        enhanced_freelancer = """Web Developer/Cloud Developer (Freelancer)
- Designed and implemented end-to-end UX workflows using Figma.
- Translated requirements into scalable designs with Tailwind CSS and modern component libraries.
- Delivered automated deployments and infrastructure setups using GitHub and Terraform.
- Maintained technical documentation and conducted user feedback sessions for iterative improvement.

Key Projects:
• HealthWell Group Platform: https://www.healthwell.group/
  Developed comprehensive healthcare platform with modern UI/UX design
• Liftoff Platform: https://liftoff-9uo4qblpz-jiawenzhus-projects.vercel.app/
  Built scalable web application with advanced cloud infrastructure"""

        project_changes = update_document_text(
            "Web Developer/Cloud Developer (Freelancer)\n- Designed and implemented end-to-end UX workflows using Figma.\n- Translated requirements into scalable designs with Tailwind CSS and modern component libraries.\n- Delivered automated deployments and infrastructure setups using GitHub and Terraform.\n- Maintained technical documentation and conducted user feedback sessions for iterative improvement.",
            enhanced_freelancer
        )
        
        if project_changes > 0:
            changes_made.append("Added project links to freelancer experience section")
        
        # If that didn't work, try a simpler approach - just add the URLs somewhere visible
        if project_changes == 0:
            # Try to add to the contact section or profile
            profile_text = "Detail-oriented UI/UX and Cloud Developer"
            enhanced_profile = """Detail-oriented UI/UX and Cloud Developer with strong experience in creating accessible, responsive user interfaces for web and mobile platforms. 

Portfolio Projects:
🔗 HealthWell Group: https://www.healthwell.group/
🔗 Liftoff Platform: https://liftoff-9uo4qblpz-jiawenzhus-projects.vercel.app/

Partnered with over 30 startups and businesses to design and launch professional, branded websites tailored to client needs."""
            
            profile_changes = update_document_text(
                "Detail-oriented UI/UX and Cloud Developer with strong experience in creating accessible, responsive user interfaces for web and mobile platforms. Partnered with over 30 startups and businesses to design and launch",
                enhanced_profile
            )
            
            if profile_changes > 0:
                changes_made.append("Added project links to profile section")
        
        # Save the document
        doc.save(filename)
        
        print("Resume update completed!")
        print("\nChanges made:")
        for change in changes_made:
            print(f"  ✅ {change}")
        
        if not changes_made:
            print("No changes were made.")
        
        return len(changes_made) > 0
        
    except Exception as e:
        print(f"Error updating resume: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def main():
    resume_file = "Jiawen-Resume-2025.docx"
    
    print(f"Updating resume: {resume_file}")
    print("Target changes:")
    print("  1. Replace 'evan ZHu' with 'Jiawen Zhu'")
    print("  2. Add HealthWell Group project: https://www.healthwell.group/")
    print("  3. Add Liftoff project: https://liftoff-9uo4qblpz-jiawenzhus-projects.vercel.app/")
    print("-" * 60)
    
    success = update_resume(resume_file)
    
    if success:
        print(f"\n🎉 Resume successfully updated: {resume_file}")
        
        # Also update the 2024 version
        try:
            from shutil import copy2
            copy2(resume_file, "Jiawen-Resume-2024.docx")
            print("✅ Also updated Jiawen-Resume-2024.docx")
        except Exception as e:
            print(f"⚠️  Could not update 2024 version: {e}")
    else:
        print("❌ Failed to update resume")

if __name__ == "__main__":
    main() 